import streamlit as st
from resume_parser import extract_text_from_resume
from match_engine import compute_match_percentage
from gpt_writer import rewrite_resume
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import re
from jd_parser import * #noqa
from recruiter_tools import (
    generate_recruiter_message,
    generate_cold_email,
    suggest_contact_titles,
    estimate_salary
)
from interview_questions import interview_questions

st.set_page_config(
    page_title="ResumePro Analytics",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Professional CSS styling - clean and corporate
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;500;600;700&display=swap');

    * {
        font-family: 'Poppins', sans-serif;
    }

    .main-header {
        background: linear-gradient(135deg, #1e293b 0%, #334155 100%);
        color: white;
        padding: 4rem 3rem;
        text-align: center;
        margin-bottom: 3rem;
        border-radius: 0 0 24px 24px;
        box-shadow: 0 8px 32px rgba(0,0,0,0.12);
    }

    .main-title {
        font-size: 2.5rem;
        font-weight: 700;
        margin-bottom: 0.8rem;
        letter-spacing: -0.02em;
    }

    .main-subtitle {
        font-size: 0.95rem;
        opacity: 0.9;
        max-width: 600px;
        margin: 0 auto;
        line-height: 1.6;
    }

    .features-container {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(280px, 1fr));
        gap: 2rem;
        margin: 4rem 0;
        padding: 0 2rem;
    }

    .feature-box {
        background: white;
        border-radius: 16px;
        padding: 2.5rem 2rem;
        box-shadow: 0 4px 24px rgba(0,0,0,0.06);
        border: 1px solid #f1f5f9;
        transition: all 0.3s ease;
        text-align: center;
    }

    .feature-box:hover {
        transform: translateY(-4px);
        box-shadow: 0 12px 40px rgba(0,0,0,0.1);
        border-color: #e2e8f0;
    }

    .feature-number {
        display: inline-block;
        width: 60px;
        height: 60px;
        background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%);
        color: white;
        border-radius: 50%;
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 1.5rem;
        font-weight: 700;
        margin-bottom: 1.5rem;
        box-shadow: 0 4px 16px rgba(59, 130, 246, 0.3);
    }

    .feature-title {
        font-size: 1.4rem;
        font-weight: 600;
        color: #1e293b;
        margin-bottom: 1rem;
    }

    .feature-description {
        color: #64748b;
        line-height: 1.6;
        font-size: 0.95rem;
    }

    .cta-section {
        background: linear-gradient(135deg, #f8fafc 0%, #f1f5f9 100%);
        border-radius: 20px;
        padding: 3rem 2rem;
        margin: 3rem 2rem;
        text-align: center;
        border: 1px solid #e2e8f0;
        box-shadow: 0 4px 16px rgba(0,0,0,0.04);
    }

    .cta-title {
        font-size: 1.6rem;
        font-weight: 600;
        color: #1e293b;
        margin-bottom: 0.8rem;
    }

    .cta-subtitle {
        color: #64748b;
        font-size: 0.9rem;
        margin-bottom: 2rem;
        max-width: 500px;
        margin-left: auto;
        margin-right: auto;
    }

    .primary-button {
        background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%);
        color: white;
        border: none;
        padding: 1.2rem 3rem;
        border-radius: 12px;
        font-size: 1.1rem;
        font-weight: 600;
        cursor: pointer;
        transition: all 0.3s ease;
        box-shadow: 0 4px 16px rgba(59, 130, 246, 0.3);
        text-decoration: none;
        display: inline-block;
    }

    .primary-button:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 24px rgba(59, 130, 246, 0.4);
    }

    .input-container {
        background: white;
        border-radius: 20px;
        padding: 3rem;
        margin: 3rem 2rem;
        box-shadow: 0 8px 32px rgba(0,0,0,0.08);
        border: 1px solid #f1f5f9;
    }

    .section-header {
        font-size: 1.8rem;
        font-weight: 600;
        color: #1e293b;
        margin-bottom: 2rem;
        text-align: center;
    }

    .form-grid {
        display: grid;
        grid-template-columns: 1fr 1fr;
        gap: 3rem;
        margin-bottom: 2rem;
    }

    .form-group {
        background: #f8fafc;
        border-radius: 12px;
        padding: 2rem;
        border: 1px solid #e2e8f0;
    }

    .form-title {
        font-size: 1.3rem;
        font-weight: 600;
        color: #1e293b;
        margin-bottom: 1.5rem;
        display: flex;
        align-items: center;
        gap: 0.5rem;
    }

    .form-title::before {
        content: '';
        width: 4px;
        height: 20px;
        background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%);
        border-radius: 2px;
    }

    .analysis-results {
        background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 100%);
        border-radius: 20px;
        padding: 3rem;
        margin: 3rem 2rem;
        border: 1px solid #bae6fd;
    }

    .results-header {
        text-align: center;
        margin-bottom: 2.5rem;
    }

    .results-title {
        font-size: 2rem;
        font-weight: 600;
        color: #1e293b;
        margin-bottom: 0.5rem;
    }

    .results-subtitle {
        color: #64748b;
        font-size: 1rem;
    }

    .metrics-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
        gap: 2rem;
        margin: 2rem 0;
    }

    .metric-card {
        background: white;
        border-radius: 16px;
        padding: 2rem;
        text-align: center;
        box-shadow: 0 4px 16px rgba(0,0,0,0.06);
        border: 1px solid #f1f5f9;
    }

    .metric-value {
        font-size: 2.5rem;
        font-weight: 700;
        color: #3b82f6;
        margin-bottom: 0.5rem;
    }

    .metric-label {
        color: #64748b;
        font-size: 0.9rem;
        text-transform: uppercase;
        letter-spacing: 0.5px;
        font-weight: 500;
    }

    .tools-section {
        background: linear-gradient(135deg, #fefefe 0%, #f8fafc 100%);
        border-radius: 20px;
        padding: 2.5rem;
        margin: 2.5rem 2rem;
        box-shadow: 0 12px 40px rgba(0,0,0,0.06);
        border: 1px solid #e2e8f0;
    }

    .tools-header {
        text-align: center;
        margin-bottom: 2rem;
    }

    .tools-title {
        font-size: 1.5rem;
        font-weight: 600;
        color: #1e293b;
        margin-bottom: 0.5rem;
    }

    .tools-subtitle {
        color: #64748b;
        font-size: 0.9rem;
    }

    .tools-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
        gap: 2rem;
    }

    .tool-card {
        background: #f8fafc;
        border-radius: 16px;
        padding: 2rem;
        border: 1px solid #e2e8f0;
        transition: all 0.3s ease;
    }

    .tool-card:hover {
        background: white;
        transform: translateY(-2px);
        box-shadow: 0 8px 24px rgba(0,0,0,0.08);
    }

    .tool-title {
        font-size: 1.3rem;
        font-weight: 600;
        color: #1e293b;
        margin-bottom: 1rem;
    }

    .tool-description {
        color: #64748b;
        line-height: 1.6;
        margin-bottom: 1.5rem;
    }

    .tool-button {
        background: linear-gradient(135deg, #10b981 0%, #059669 100%);
        color: white;
        border: none;
        padding: 0.8rem 1.5rem;
        border-radius: 8px;
        font-size: 0.95rem;
        font-weight: 500;
        cursor: pointer;
        transition: all 0.2s ease;
        width: 100%;
    }

    .tool-button:hover {
        transform: translateY(-1px);
        box-shadow: 0 4px 12px rgba(16, 185, 129, 0.3);
    }

    .footer {
        background: #1e293b;
        color: white;
        padding: 2rem 1.5rem;
        text-align: center;
        margin-top: 3rem;
        border-radius: 12px 12px 0 0;
    }

    .footer-title {
        font-size: 1.2rem;
        font-weight: 600;
        margin-bottom: 0.3rem;
    }

    .footer-subtitle {
        opacity: 0.8;
        margin-bottom: 0.8rem;
        font-size: 0.9rem;
    }

    .footer-note {
        font-size: 0.75rem;
        opacity: 0.6;
    }

    .status-message {
        padding: 1rem 1.5rem;
        border-radius: 8px;
        margin: 1rem 0;
        font-weight: 500;
    }

    .status-success {
        background: #dcfce7;
        border: 1px solid #bbf7d0;
        color: #166534;
    }

    .status-info {
        background: #dbeafe;
        border: 1px solid #bfdbfe;
        color: #1e40af;
    }

    .status-warning {
        background: #fef3c7;
        border: 1px solid #fde68a;
        color: #92400e;
    }
</style>
""", unsafe_allow_html=True)

# Main Header
st.markdown("""
<div class="main-header">
    <div class="main-title">ResumePro Analytics</div>
    <div class="main-subtitle">
        Professional AI-powered resume analysis and optimization platform for career advancement
    </div>
</div>
""", unsafe_allow_html=True)

# Features Section
st.markdown("""
<div class="features-container">
    <div class="feature-box">
        <div class="feature-number">1</div>
        <div class="feature-title">Intelligent Analysis</div>
        <div class="feature-description">Advanced algorithms evaluate resume-job compatibility with detailed scoring and insights.</div>
    </div>
    <div class="feature-box">
        <div class="feature-number">2</div>
        <div class="feature-title">AI Optimization</div>
        <div class="feature-description">Automatically rewrite and enhance resumes to better match job requirements and ATS systems.</div>
    </div>
    <div class="feature-box">
        <div class="feature-number">3</div>
        <div class="feature-title">Interview Preparation</div>
        <div class="feature-description">Generate personalized interview questions and receive expert preparation guidance.</div>
    </div>
    <div class="feature-box">
        <div class="feature-number">4</div>
        <div class="feature-title">Career Tools</div>
        <div class="feature-description">Access professional networking templates, salary insights, and career development resources.</div>
    </div>
</div>
""", unsafe_allow_html=True)


# ---- Export Helpers ----

def _parse_markdown_line(line):
    """Determine the type and content of a markdown line."""
    stripped = line.strip()
    
    if stripped.startswith("### "):
        return "h3", stripped[4:].strip()
    elif stripped.startswith("## "):
        return "h2", stripped[3:].strip()
    elif stripped.startswith("# "):
        return "h1", stripped[2:].strip()
    elif stripped.startswith("- ") or stripped.startswith("* "):
        return "bullet", stripped[2:].strip()
    elif stripped == "---" or stripped == "***":
        return "separator", ""
    elif stripped == "":
        return "blank", ""
    else:
        return "text", stripped


def _add_bold_runs(paragraph, text):
    """Parse **bold** markers and add runs with appropriate formatting."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def download_resume_docx(resume_text):
    """Convert markdown-formatted resume text into a properly formatted Word document."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Set narrow margins for resume
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    lines = resume_text.split('\n')
    
    for line in lines:
        line_type, content = _parse_markdown_line(line)
        
        if line_type == "blank":
            continue  # Skip blank lines to keep resume tight
            
        elif line_type == "h1":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_bold_runs(p, content)
            for run in p.runs:
                run.font.size = Pt(16)
                run.bold = True
            p.space_after = Pt(4)
            p.space_before = Pt(0)
            
        elif line_type == "h2":
            p = doc.add_paragraph()
            _add_bold_runs(p, content.upper())
            for run in p.runs:
                run.font.size = Pt(12)
                run.bold = True
            p.space_after = Pt(2)
            p.space_before = Pt(8)
            # Add a bottom border to section headers
            from docx.oxml.ns import qn
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): '4',
                qn('w:space'): '1',
                qn('w:color'): '000000'
            })
            pBdr.append(bottom)
            pPr.append(pBdr)

        elif line_type == "h3":
            p = doc.add_paragraph()
            _add_bold_runs(p, content)
            for run in p.runs:
                run.font.size = Pt(11)
                run.bold = True
            p.space_after = Pt(1)
            p.space_before = Pt(4)

        elif line_type == "bullet":
            p = doc.add_paragraph(style='List Bullet')
            _add_bold_runs(p, content)
            for run in p.runs:
                run.font.size = Pt(10.5)
            p.space_after = Pt(1)
            p.space_before = Pt(0)
            p.paragraph_format.left_indent = Inches(0.25)
            
        elif line_type == "separator":
            continue  # Skip markdown separators
            
        else:  # regular text
            p = doc.add_paragraph()
            _add_bold_runs(p, content)
            for run in p.runs:
                run.font.size = Pt(10.5)
            p.space_after = Pt(2)
            p.space_before = Pt(0)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def download_resume_pdf(resume_text):
    """Convert markdown-formatted resume text into a properly formatted PDF."""
    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
        leftMargin=0.6 * inch,
        rightMargin=0.6 * inch
    )

    styles = getSampleStyleSheet()

    # Custom styles for resume sections
    styles.add(ParagraphStyle(
        name='ResumeName',
        parent=styles['Title'],
        fontSize=16,
        leading=20,
        spaceAfter=4,
        spaceBefore=0,
        alignment=1  # center
    ))

    styles.add(ParagraphStyle(
        name='SectionHeader',
        parent=styles['Heading2'],
        fontSize=12,
        leading=16,
        spaceAfter=4,
        spaceBefore=10,
        textColor='black',
        borderWidth=0.5,
        borderPadding=2,
        borderColor='black',
    ))

    styles.add(ParagraphStyle(
        name='SubHeader',
        parent=styles['Heading3'],
        fontSize=11,
        leading=14,
        spaceAfter=2,
        spaceBefore=6,
        textColor='black',
    ))

    styles.add(ParagraphStyle(
        name='BulletItem',
        parent=styles['Normal'],
        fontSize=10,
        leading=13,
        spaceAfter=2,
        spaceBefore=0,
        leftIndent=18,
        bulletIndent=6,
    ))

    styles.add(ParagraphStyle(
        name='ResumeText',
        parent=styles['Normal'],
        fontSize=10,
        leading=13,
        spaceAfter=2,
        spaceBefore=0,
    ))

    story = []
    lines = resume_text.split('\n')

    for line in lines:
        line_type, content = _parse_markdown_line(line)

        # Convert markdown bold to reportlab bold tags
        content = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', content)

        if line_type == "blank":
            continue

        elif line_type == "h1":
            story.append(Paragraph(content, styles['ResumeName']))

        elif line_type == "h2":
            # Add a thin line above section headers
            story.append(Spacer(1, 4))
            story.append(Paragraph(f"<u><b>{content.upper()}</b></u>", styles['SectionHeader']))

        elif line_type == "h3":
            story.append(Paragraph(f"<b>{content}</b>", styles['SubHeader']))

        elif line_type == "bullet":
            story.append(Paragraph(f"&bull;  {content}", styles['BulletItem']))

        elif line_type == "separator":
            continue

        else:
            story.append(Paragraph(content, styles['ResumeText']))

    doc.build(story)
    buffer.seek(0)
    return buffer


def download_resume_txt(resume_text):
    """Download resume as plain text."""
    return resume_text.encode('utf-8')


# Call-to-Action Section
st.markdown("""
<div class="cta-section">
    <div class="cta-title">Get Started with AI Analysis</div>
    <div class="cta-subtitle">Upload your resume and target job description to unlock professional insights and optimization tools</div>
</div>
""", unsafe_allow_html=True)

# Input Section
st.markdown('<div class="input-container">', unsafe_allow_html=True)
st.markdown('<div class="section-header">Document Input</div>', unsafe_allow_html=True)

st.markdown('<div class="form-grid">', unsafe_allow_html=True)

# Resume Input
st.markdown("""
<div class="form-group">
    <div class="form-title">Resume Upload</div>
</div>
""", unsafe_allow_html=True)

resume_input_method = st.radio(
    "Select input method:",
    ("File Upload", "Text Input"),
    key="resume_input",
    horizontal=True
)

resume_text = ""
if resume_input_method == "File Upload":
    resume_file = st.file_uploader(
        "Choose resume file (.pdf or .docx)",
        type=["pdf", "docx"],
        help="Maximum file size: 200MB"
    )

    if resume_file:
        with st.spinner("Extracting resume content..."):
            resume_text = extract_text_from_resume(resume_file).strip()
        st.markdown('<div class="status-message status-success">Resume processed successfully</div>', unsafe_allow_html=True)

else:
    resume_text = st.text_area(
        "Paste resume content:",
        height=250,
        placeholder="Copy and paste your complete resume text here..."
    )

# Job Description Input
st.markdown("""
<div class="form-group">
    <div class="form-title">Job Description</div>
</div>
""", unsafe_allow_html=True)

job_description = st.text_area(
    "Paste job description:",
    height=250,
    placeholder="Copy and paste the complete job description you want to analyze against..."
)

if job_description.strip():
    # Clear cached job info if JD text changed so extraction re-runs
    if "last_jd_hash" not in st.session_state or st.session_state.last_jd_hash != hash(job_description):
        st.session_state.last_jd_hash = hash(job_description)
        st.session_state.jd_info_extracted = False

    with st.spinner("Analyzing job requirements..."):
        job_description_parsed = parse_jd(job_description)
        job_description = job_description_parsed["clean_text"]
    st.markdown('<div class="status-message status-success">Job description analyzed successfully</div>', unsafe_allow_html=True)
else:
    job_description = ""

st.markdown('</div>', unsafe_allow_html=True)  # End form-grid
st.markdown('</div>', unsafe_allow_html=True)  # End input-container

# Analysis Section
if resume_text and job_description:
    st.markdown('<div class="analysis-results">', unsafe_allow_html=True)
    st.markdown("""
    <div class="results-header">
        <div class="results-title">AI Analysis Results</div>
        <div class="results-subtitle">Professional insights powered by advanced machine learning</div>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns(2)

    with col1:
        if st.button("Analyze Compatibility", key="analyze_btn", help="Get detailed compatibility score between your resume and job description"):
            with st.spinner("AI analyzing compatibility..."):
                try:
                    match_percent = compute_match_percentage(resume_text, job_description)

                    # Main score display
                    st.markdown(f"""
                    <div style="background: linear-gradient(135deg, #10b981 0%, #059669 100%); color: white; padding: 2.5rem; border-radius: 20px; text-align: center; margin: 1.5rem 0; box-shadow: 0 12px 32px rgba(16, 185, 129, 0.3);">
                        <div style="font-size: 3.5rem; font-weight: 700; margin-bottom: 0.5rem;">{match_percent}%</div>
                        <div style="font-size: 1.3rem; opacity: 0.9; font-weight: 500;">Compatibility Score</div>
                        <div style="font-size: 0.95rem; opacity: 0.8; margin-top: 0.5rem;">Resume-Job Match</div>
                    </div>
                    """, unsafe_allow_html=True)

                    # Detailed metrics
                    resume_words = len(resume_text.split())
                    jd_words = len(job_description.split())

                    st.markdown('<div class="metrics-grid">', unsafe_allow_html=True)
                    st.markdown(f"""
                    <div class="metric-card">
                        <div class="metric-value">{resume_words}</div>
                        <div class="metric-label">Resume Words</div>
                    </div>
                    <div class="metric-card">
                        <div class="metric-value">{jd_words}</div>
                        <div class="metric-label">Job Description Words</div>
                    </div>
                    <div class="metric-card">
                        <div class="metric-value">{match_percent}%</div>
                        <div class="metric-label">Match Score</div>
                    </div>
                    """, unsafe_allow_html=True)
                    st.markdown('</div>', unsafe_allow_html=True)

                except Exception as e:
                    st.error(f"Analysis failed: {str(e)}")

    with col2:
        if st.button("Optimize Resume", key="optimize_btn", help="AI-powered resume enhancement for better job matching"):
            with st.spinner("AI optimizing your resume..."):
                try:
                    edited_resume = rewrite_resume(resume_text, job_description)
                    st.session_state["edited_resume"] = edited_resume

                    st.markdown("""
                    <div style="background: linear-gradient(135deg, #8b5cf6 0%, #7c3aed 100%); color: white; padding: 2.5rem; border-radius: 20px; text-align: center; margin: 1.5rem 0; box-shadow: 0 12px 32px rgba(139, 92, 246, 0.3);">
                        <div style="font-size: 3rem; font-weight: 700; margin-bottom: 1rem;">✓</div>
                        <div style="font-size: 1.5rem; font-weight: 600; margin-bottom: 0.5rem;">Resume Optimized</div>
                        <div style="font-size: 1rem; opacity: 0.9;">Enhanced for better ATS compatibility and job matching</div>
                    </div>
                    """, unsafe_allow_html=True)

                except Exception as e:
                    st.error(f"Optimization failed: {str(e)}")

    # Download section
    if "edited_resume" in st.session_state:
        st.markdown("""
        <div style="background: white; border-radius: 20px; padding: 3rem; margin: 3rem 0; box-shadow: 0 8px 32px rgba(0,0,0,0.08); border: 1px solid #f1f5f9;">
            <h3 style="color: #1e293b; margin-bottom: 2rem; text-align: center; font-size: 1.8rem;">Download Enhanced Resume</h3>
            <p style="text-align: center; color: #64748b; margin-bottom: 2rem;">Choose your preferred format for the optimized resume</p>
        </div>
        """, unsafe_allow_html=True)

        col1, col2, col3 = st.columns(3)

        with col1:
            docx_file = download_resume_docx(st.session_state["edited_resume"])
            st.download_button(
                label="Word Document",
                data=docx_file,
                file_name="ResumePro_Optimized.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_docx",
                help="Download as Microsoft Word document"
            )

        with col2:
            pdf_file = download_resume_pdf(st.session_state["edited_resume"])
            st.download_button(
                label="PDF Document",
                data=pdf_file,
                file_name="ResumePro_Optimized.pdf",
                mime="application/pdf",
                key="download_pdf",
                help="Download as PDF document"
            )

        with col3:
            txt_file = download_resume_txt(st.session_state["edited_resume"])
            st.download_button(
                label="Plain Text",
                data=txt_file,
                file_name="ResumePro_Optimized.txt",
                mime="text/plain",
                key="download_txt",
                help="Download as plain text file"
            )

    st.markdown('</div>', unsafe_allow_html=True)

else:
    st.markdown("""
    <div class="status-message status-warning" style="text-align: center; font-size: 1.1rem; margin: 3rem 0;">
        <strong>Ready for Professional Analysis?</strong><br>
        Please provide both your resume and target job description above to access AI-powered insights and optimization tools.
    </div>
    """, unsafe_allow_html=True)

# Professional Tools Section
if resume_text and job_description:
    st.markdown('<div class="tools-section">', unsafe_allow_html=True)
    st.markdown("""
    <div class="tools-header">
        <div class="tools-title">Professional Career Tools</div>
        <div class="tools-subtitle">Advanced AI-powered resources for career advancement</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div class="tools-grid">', unsafe_allow_html=True)

    # Networking & Communication Tools
    st.markdown("""
    <div class="tool-card">
        <div class="tool-title">Networking & Communication</div>
        <div class="tool-description">Generate professional LinkedIn messages and email templates to connect with recruiters and hiring managers.</div>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Generate LinkedIn Message", key="linkedin_btn", help="Create a professional LinkedIn connection message"):
            with st.spinner("Crafting message..."):
                try:
                    recruiter_msg = generate_recruiter_message(job_description)
                    st.markdown("""
                    <div style="background: #f0f9ff; border: 1px solid #3b82f6; border-radius: 8px; padding: 1rem; margin: 1rem 0;">
                        <h4 style="margin: 0 0 1rem 0; color: #1e40af;">Professional LinkedIn Message:</h4>
                    </div>
                    """, unsafe_allow_html=True)
                    st.text_area("", recruiter_msg, height=120, key="linkedin_msg")
                except Exception as e:
                    st.error(f"Message generation failed: {str(e)}")

    with col2:
        if st.button("Generate Cold Email", key="email_btn", help="Create a professional cold outreach email"):
            with st.spinner("Creating email template..."):
                try:
                    cold_email = generate_cold_email(job_description)
                    st.markdown("""
                    <div style="background: #f0f9ff; border: 1px solid #3b82f6; border-radius: 8px; padding: 1rem; margin: 1rem 0;">
                        <h4 style="margin: 0 0 1rem 0; color: #1e40af;">Professional Email Template:</h4>
                    </div>
                    """, unsafe_allow_html=True)
                    st.text_area("", cold_email, height=120, key="email_msg")
                except Exception as e:
                    st.error(f"Email generation failed: {str(e)}")

    # Contact Strategy
    if st.button("Contact Strategy", key="contacts_btn", help="Get suggestions for key people to reach out to"):
        with st.spinner("Analyzing contact opportunities..."):
            try:
                titles = suggest_contact_titles(job_description)
                st.markdown("""
                <div style="background: #f0fdf4; border: 1px solid #22c55e; border-radius: 8px; padding: 1rem; margin: 1rem 0;">
                    <h4 style="margin: 0 0 1rem 0; color: #15803d;">Strategic Contact Suggestions:</h4>
                </div>
                """, unsafe_allow_html=True)
                st.text_area("", ", ".join(titles), height=80, key="contacts_list")
            except Exception as e:
                st.error(f"Contact analysis failed: {str(e)}")

    # Salary Intelligence
    if st.button("Salary Intelligence", key="salary_btn", help="Get market salary insights for this role"):
        with st.spinner("Analyzing compensation data..."):
            try:
                salary = estimate_salary(job_description)
                st.markdown(f"""
                <div style="background: linear-gradient(135deg, #fbbf24 0%, #f59e0b 100%); color: white; padding: 2rem; border-radius: 16px; text-align: center; margin: 1rem 0; box-shadow: 0 8px 25px rgba(245, 158, 11, 0.3);">
                    <div style="font-size: 1.8rem; font-weight: 600; margin-bottom: 0.5rem;">Salary Range Estimate</div>
                    <div style="font-size: 1.2rem; opacity: 0.9;">{salary}</div>
                </div>
                """, unsafe_allow_html=True)
            except Exception as e:
                st.error(f"Salary analysis failed: {str(e)}")

    st.markdown('</div>', unsafe_allow_html=True)  # End tools grid
    st.markdown('</div>', unsafe_allow_html=True)  # End tools section

    # Interview Preparation Section
    st.markdown("""
    <div style="background: linear-gradient(135deg, #a855f7 0%, #7c3aed 100%); color: white; padding: 2rem 1.5rem; border-radius: 16px; text-align: center; margin: 2rem 2rem; box-shadow: 0 8px 24px rgba(168, 85, 247, 0.2);">
        <div style="font-size: 1.4rem; font-weight: 600; margin-bottom: 0.5rem;">Interview Preparation</div>
        <div style="font-size: 0.9rem; opacity: 0.9;">AI-generated questions tailored to your profile and target role</div>
    </div>
    """, unsafe_allow_html=True)

    if st.button("Generate Interview Questions", key="interview_btn", help="Get personalized interview questions for this role"):
        with st.spinner("Creating interview questions..."):
            try:
                questions = interview_questions(resume_text, job_description)
                st.markdown("""
                <div style="background: white; border-radius: 12px; padding: 1.5rem; margin: 1.5rem 2rem; box-shadow: 0 4px 16px rgba(0,0,0,0.06); border: 1px solid #f1f5f9;">
                    <h4 style="color: #1e293b; margin-bottom: 1rem; text-align: center;">Your Interview Preparation Questions</h4>
                </div>
                """, unsafe_allow_html=True)
                st.text_area("", questions, height=350, key="interview_questions")
            except Exception as e:
                st.error(f"Interview question generation failed: {str(e)}")

    # Career Development Section
    st.markdown("""
    <div style="background: linear-gradient(135deg, #06b6d4 0%, #0891b2 100%); color: white; padding: 2rem 1.5rem; border-radius: 16px; text-align: center; margin: 2rem 2rem; box-shadow: 0 8px 24px rgba(6, 182, 212, 0.2);">
        <div style="font-size: 1.4rem; font-weight: 600; margin-bottom: 0.5rem;">Career Development</div>
        <div style="font-size: 0.9rem; opacity: 0.9;">Expert tips and strategies for career advancement</div>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns(2)

    with col1:
        if st.button("Resume Enhancement Tips", key="tips_btn", help="Get personalized resume improvement suggestions"):
            with st.spinner("Analyzing your resume..."):
                try:
                    tips = []
                    if len(resume_text.split()) < 200:
                        tips.append("Consider adding more details about your experience and achievements.")
                    if "education" not in resume_text.lower():
                        tips.append("Add your educational background if relevant to the position.")
                    if len([line for line in resume_text.split('\n') if line.strip().startswith('-')]) < 5:
                        tips.append("Include more bullet points with specific accomplishments and metrics.")

                    if not tips:
                        tips = ["Your resume appears comprehensive and well-structured."]

                    st.markdown("""
                    <div style="background: #f0fdf4; border: 1px solid #22c55e; border-radius: 8px; padding: 1.5rem; margin: 1rem 2rem;">
                        <h4 style="margin: 0 0 1rem 0; color: #15803d;">Resume Enhancement Recommendations:</h4>
                    </div>
                    """, unsafe_allow_html=True)
                    for tip in tips:
                        st.markdown(f"• {tip}")
                except Exception as e:
                    st.error(f"Resume analysis failed: {str(e)}")

    with col2:
        if st.button("Job Search Strategy", key="strategy_btn", help="Get proven job search best practices"):
            st.markdown("""
            <div style="background: #fef3c7; border: 1px solid #f59e0b; border-radius: 8px; padding: 1.5rem; margin: 1rem 2rem;">
                <h4 style="margin: 0 0 1rem 0; color: #92400e;">Proven Job Search Strategies:</h4>
            </div>
            """, unsafe_allow_html=True)
            st.markdown("• **Personalize** each application with company-specific details")
            st.markdown("• **Network** actively on professional platforms and industry events")
            st.markdown("• **Follow up** on applications within 1-2 weeks")
            st.markdown("• **Research** thoroughly before interviews and networking calls")
            st.markdown("• **Track** your progress and maintain detailed application records")

else:
    st.markdown("""
    <div style="background: #f8fafc; border: 1px solid #e2e8f0; color: #64748b; padding: 3rem; border-radius: 16px; text-align: center; margin: 3rem 2rem;">
        <div style="font-size: 3rem; margin-bottom: 1rem;">🔒</div>
        <div style="font-size: 1.5rem; font-weight: 600; margin-bottom: 0.5rem;">Professional Tools Unlocked</div>
        <div style="font-size: 1rem;">Provide your resume and job description above to access advanced career development tools and AI-powered insights.</div>
    </div>
    """, unsafe_allow_html=True)

# Footer
st.markdown("""
<div class="footer">
    <div class="footer-title">ResumePro Analytics</div>
    <div class="footer-subtitle">Professional AI-powered resume analysis and optimization platform</div>
    <div class="footer-note">Powered by Google Gemini AI • Built with Streamlit</div>
</div>
""", unsafe_allow_html=True)
